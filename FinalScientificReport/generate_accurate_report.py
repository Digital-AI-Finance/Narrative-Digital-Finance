"""Generate accurate SNSF Scientific Report with verified content and character counts."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Create new document
doc = Document()

# Title
title = doc.add_heading('Scientific Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('SNSF Grant IZCOZ0_213370 - Narrative Digital Finance')
doc.add_paragraph('Principal Investigator: Prof. Dr. Joerg Osterrieder')
doc.add_paragraph('Institution: Bern University of Applied Sciences / University of Twente')
doc.add_paragraph('Period: July 2023 - June 2026')
doc.add_paragraph()

# Section 1.1: Achievement of Research Objectives (max 6000 chars)
doc.add_heading('1.1 Achievement of Research Objectives', level=1)

section_1_1 = """The project successfully achieved all research objectives across the four work packages defined in the research plan, producing significant methodological and empirical contributions to the field of narrative digital finance.

WP1: Text Data & Text Analytics
We developed comprehensive NLP frameworks for financial text analysis. The TOPol (Topic-Orientation Polarity) framework provides a semi-unsupervised framework for reconstructing multidimensional semantic polarity fields using transformer embeddings, UMAP projection, and Leiden clustering. This moves beyond traditional unidimensional sentiment scoring to capture the full complexity of financial narratives. TOPol begins by embedding documents using a general-purpose transformer-based LLM, followed by neighbor-tuned UMAP projection and topic-based segmentation via Leiden partitioning. Given a contextual boundary between regimes A and B, the framework computes directional vectors between corresponding topic-boundary centroids, producing a polarity field that captures fine-grained semantic displacement. The framework generates contrastive labels with estimated coverage to interpret identified polarity shifts. Code and data are publicly available at OSF (osf.io/nr94j). Additionally, we created an AI-enhanced systematic literature review methodology using NLP techniques, clustering algorithms, and interpretability tools, currently under revision at Financial Innovation journal. This SLR applies algorithmic selection to Scopus peer-reviewed literature, identifying research practices for modeling financial narratives with NLP methods.

WP2: Structural Breaks Detection & Asset Price Bubbles
We developed methodologies for detecting structural breaks in financial markets. The macro narratives framework combines a PCA-based US Macro Strength Index (using FED Funds Rate, CPI, PPI, GDP, Unemployment, Nonfarm Payrolls from St. Louis FED FRED, 1996-2025) with the PELT algorithm for change point detection. The first principal component explains the largest share of variance and captures joint dynamics in growth, inflation, and labor conditions. As a case study, the framework was applied to Federal Reserve banker speech transcripts (BIS Gigando datasets, 2004-2010) around the 2008 financial crisis, successfully identifying narrative shifts at the detected breakpoint (2007-05-01) using BERTopic. Null hypothesis testing via 1000 random shuffling simulations validated statistical significance of observed polarity drift, semantic centroid movement, and keyword evolution. Related prework on NFT bubble detection (Wang et al., 2022, Journal of Chinese Economic and Business Studies) has been extended for PhD research.

WP3: Narratives for Structural Breaks
We established foundational frameworks for understanding how narratives shape financial markets. The Multimodal Influence paper (SSRN 4698153) extends traditional NLP to include non-textual data (images, video, audio) for financial instrument pricing using Generative AI. This research acknowledges the exponential growth from a few dozen exabytes in 2003 to approximately 180 zettabytes by 2025, requiring multimodal analysis approaches. The BERTopic-based narrative shift analysis quantifies polarity drift (FinBERT sentiment), semantic centroid movement (cosine similarity between topic centroids), and keyword evolution (Jaccard similarity of Maximal Marginal Relevance keyword sets) across economic regimes. Word clouds aggregate MMR keywords from all topics before and after breakpoints, visualizing thematic drift across narrative regimes.

WP4: Multidimensional AI and ML Solutions
In collaboration with Deutsche Borse, we developed novel methodology for classifying market participants based on nanosecond-level reaction times. The HFT Market Microstructure framework uses nanosecond-level timestamp data from Deutsche Borse's T7 platform to distinguish Ultra-Fast Traders (UFT, <1 microsecond), High-Frequency Traders (HFT, 1-10 microseconds), and conventional participants (>10 microseconds). UFTs rely on Field-Programmable Gate Arrays (FPGAs) for ultra-low-latency execution. Research documents participation shares (UFT/HFT produce roughly 20% of total participation), price discovery via 15-second mark-out signal-to-noise ratio, and market quality metrics including order-imbalance volatility, Amihud illiquidity, and high-frequency return diagnostics (autocorrelation, variance-ratio tests) for Euro STOXX 50 Index Futures (FESX) from January to August 2025.

Key Publications: TOPol Semantic Polarity Fields (Taibi, Gomez, 2025); SLR on Financial Narratives (Taibi et al., 2025, under revision at Financial Innovation); Nanosecond Microstructure: HFT Participation Stylized Facts (Taibi, Osterrieder, Schlamp, 2025); Reaction Times to Economic News in HFT (Osterrieder, Schlamp, 2025, SSRN 5112295); Macro Narratives & Fed Communication (Taibi, 2025, conference poster); Multimodal Influence (Bolesta, Taibi et al., 2024, SSRN 4698153).

Data Infrastructure: Custom Python pipelines for data ingestion from RavenPack (news headlines), LSEG (earnings call transcripts), BIS (central bank speeches), SEC EDGAR (10-K/10-Q filings), and Deutsche Borse (nanosecond-level Xetra/Eurex trading data). The realized-library provides efficient Python and C++ implementations for realized volatility estimators and jump detection."""

print(f"Section 1.1 character count: {len(section_1_1)} (limit: 6000)")
doc.add_paragraph(section_1_1)

# Section 1.2: Challenges (max 2000 chars)
doc.add_heading('1.2 Challenges, Negative Results and Unexpected Outcomes', level=1)

section_1_2 = """Data Access and Infrastructure
The project developed custom data pipelines for RavenPack financial news, LSEG earnings transcripts, BIS central bank speeches, SEC EDGAR filings (10-K, 10-Q), and Deutsche Borse high-frequency trading data. While the original proposal envisioned using Bloomberg and Refinitiv terminals, these alternative sources proved more suitable for the research objectives, providing both textual narratives and nanosecond-level market reactions.

Methodological Evolution
The original WP3.1 design planned controlled 2x2 investment experiments to test narrative effects. This approach evolved toward the multimodal influence framework (SSRN 4698153), providing a more comprehensive theoretical foundation for understanding how narratives across multiple modalities (text, images, video, audio) affect financial instrument pricing. The TOPol framework similarly moved beyond unidimensional sentiment scoring to capture multidimensional semantic polarity fields through transformer embeddings and Leiden clustering.

Unexpected Opportunity: Deutsche Borse Collaboration
Access to nanosecond-level trading data from Deutsche Borse through Dr. Stefan Schlamp (Head of Quantitative Analytics) enabled complementary HFT market microstructure research, analyzing Euro STOXX 50 Index Futures from the T7 platform. This collaboration began January 2024 and produced joint publications on UFT/HFT classification methodology."""

print(f"Section 1.2 character count: {len(section_1_2)} (limit: 2000)")
doc.add_paragraph(section_1_2)

# Section 1.3: Contribution to Knowledge (max 2000 chars)
doc.add_heading('1.3 Contribution to Knowledge Advancement', level=1)

section_1_3 = """This research bridges narrative economics (Shiller, 2017, 2019) with computational linguistics and financial econometrics, providing quantitative tools for analyzing how stories and interpretations influence market dynamics and volatility regimes. The work addresses a fundamental question: how do forms of narratives influence volatility and regime changes in financial markets?

TOPol Framework: First semi-unsupervised framework for multidimensional semantic polarity in financial texts, moving beyond unidimensional sentiment scoring to capture narrative complexity through transformer embeddings, UMAP projection, and Leiden clustering. The vectorial representation enables quantification of the magnitude, direction, and semantic meaning of polarity shifts across discourse regime changes.

Macro-Narrative Integration: Novel combination of PCA-based macroeconomic indices (US Macro Strength Index from FRED data) with BERTopic topic modeling for detecting narrative regime shifts. Null hypothesis testing via 1000 random shuffling simulations validates statistical significance.

HFT Classification: New methodology using nanosecond latency thresholds to distinguish UFT (<1 microsecond, FPGA-based), HFT (1-10 microseconds), and conventional market participants (>10 microseconds), with documented participation shares, price discovery metrics, and market quality indicators.

SLR Methodology: AI-enhanced systematic literature review framework using NLP techniques, clustering algorithms, and interpretability tools for reproducible, efficient research synthesis.

Impact Areas: Narrative Economics, Central Bank Communication Analysis, Market Microstructure, Volatility Regime Detection, NLP in Finance, High-Frequency Trading Research, Asset Bubble Detection."""

print(f"Section 1.3 character count: {len(section_1_3)} (limit: 2000)")
doc.add_paragraph(section_1_3)

# Section 1.4: Collaboration (max 4000 chars)
doc.add_heading('1.4 Additional Reporting Requirement', level=1)

section_1_4a = """Research Collaborations

Deutsche Borse AG: Dr. Stefan Schlamp, Head of Quantitative Analytics, Market Data and Services. Provided nanosecond-level Xetra/Eurex trading data from the T7 platform for Euro STOXX 50 Index Futures and DAX Futures, along with MSCI World and S&P500 iShares ETFs. Joint research on HFT market microstructure produced novel methodology for classifying Ultra-Fast Traders (UFT, <1 microsecond using FPGAs), High-Frequency Traders (HFT, 1-10 microseconds), and conventional market participants. Collaboration began January 2024, producing the paper "Nanosecond Microstructure: High-Frequency Traders Participation Stylized Facts" (Taibi, Osterrieder, Schlamp, 2025) and "Reaction Times to Economic News in High-Frequency Trading" (Osterrieder, Schlamp, 2025, SSRN 5112295).

Quoniam Asset Management (Germany): Dr. Axel Gross-Klussmann supervised PhD internship on quantitative investment strategies, connecting narrative research with practical asset management applications. This industry collaboration bridges academic research with real-world portfolio management and provided guidance on practical applications of narrative-based trading signals.

University of Twente: Joint PhD supervision and dual affiliation. Gabin Taibi's doctoral research on "Modeling Narrative Dynamics for Volatility Regime Detection" directly contributes to project objectives (December 2023 - November 2027). Promotor: Prof. Dr. Joerg Osterrieder; Co-Promotor: Dr. Xiaohong Huang; Supervisors: Dr. Stefan Schlamp, Dr. Axel Gross-Klussmann. The research addresses how narrative information forms a latent context for market decisions.

COST Action CA19130 Academic Collaborators: Prof. Dr. Wolfgang Karl Haerdle (Humboldt-University Berlin, Ladislaus von Bortkiewicz Professor of Statistics), Prof. Dr. Daniel Traian Pele (Bucharest University of Economic Studies, Department of Statistics and Econometrics, ORCID: 0000-0002-5891-5495), and Prof. Dr. Codruta Mare (Babes-Bolyai University) contributed through joint research activities, PhD training, and the COST FinAI network. The Multimodal Influence paper (SSRN 4698153) was co-authored with Prof. Mare and Prof. Hopp.

PhD Research Progress: The research progresses through five stages: (1) Data Collection & SLR (completed - paper under revision at Financial Innovation), (2) Narrative Detection (in progress - TOPol framework published on OSF), (3) Volatility Estimation (2026 - realized-library for volatility estimators, Hurst exponent computation, and jump detection), (4) Narrative-Volatility Integration using Bai-Perron, CUSUM, and Bayesian frameworks to link regime shifts with narrative dynamics (2026-2027), and (5) Thesis Synthesis (2027). The PhD qualification report was successfully completed in November 2025 with examination by Prof. Dr. Wolfgang Haerdle, Prof. Dr. Ali Hirsa (Columbia University), and Prof. Dr. Daniel Pele.

Prework on Asset Bubbles: Related prework on NFT bubble detection (Wang, Horky, Baals, Lucey, Vigne, 2022, Journal of Chinese Economic and Business Studies) has been extended and updated for ongoing PhD research, directly addressing WP2 objectives on asset price bubble detection. This work applies econometric bubble detection methods to cryptocurrency and NFT markets."""

print(f"Section 1.4a (Collaboration) character count: {len(section_1_4a)} (limit: 4000)")
doc.add_paragraph(section_1_4a)

# Section 1.4b: COST Action (max 4000 chars)
section_1_4b = """COST Action CA19130 - Fintech and AI in Finance

As Chair of COST Action CA19130 (2020-2024, successfully concluded October 2024), the project benefited from extensive pan-European collaboration: 49 participating countries, 270+ researchers, EU H2020 & Horizon Europe funded. The COST network provided a platform for disseminating project results and recruiting collaborators across European institutions.

Added Value from COST Participation:

Annual Conferences: Organized 7th and 8th European COST Conferences on AI in Finance at Bern University of Applied Sciences (September 2023, September 2024). These conferences brought together leading researchers and practitioners in AI-driven finance, providing a platform for presenting project results on narrative analysis, NLP in finance, and market microstructure research. The 9th conference is planned for 2025 as post-COST continuation.

PhD Training: COST FinAI PhD School 2024 in Treviso, Italy focused on Generative AI, Chatbots, and LLMs in Finance. This training directly benefited project PhD researcher Gabin Taibi and provided networking opportunities with 30+ doctoral students from across Europe working on AI in finance. Additional PhD training occurred at the 16th ERCIM Working Group Conference (December 2023, Berlin).

Working Groups: Active coordination across 6 Working Groups covering AI applications in finance, blockchain technologies, regulatory technology (RegTech), sustainable finance, and market microstructure. Project research contributes to WG1 (AI in Finance) and WG6 (Market Microstructure) through narrative analysis and HFT research.

Joint Publications: The COST network produced collaborative research including "Mitigating Digital Asset Risks" (Teng, Härdle, Osterrieder, Baals et al., 2023, SSRN 4594467) with 30+ co-authors from COST member institutions, addressing blockchain, DLT, and digital asset regulation.

International Meetings: COST FinAI Meets Istanbul (May 2024) and Brussels (2024) facilitated researcher mobility and knowledge exchange. These events enabled discussions with European Central Bank researchers and regulatory bodies on the practical implications of narrative analysis for financial stability. A keynote presentation was delivered at American University of Sharjah (May 2024) on Data Science in Finance.

Knowledge Transfer: Research stays and personnel exchange with industry partners including Deutsche Borse and Quoniam Asset Management. These exchanges ensured that academic research addresses real-world challenges in quantitative finance and market monitoring. The collaboration produced methodologies applicable to regulatory monitoring and market surveillance.

MSCA Industrial Doctoral Network: As Coordinator of the EU Horizon Europe MSCA Industrial Doctoral Network on Digital Finance (Grant Agreement No. 101119635, 4.5M EUR, 13 institutions, 100+ researchers), the project contributes to training the next generation of researchers at the intersection of AI, NLP, and quantitative finance. This network extends the project's impact beyond the SNSF funding period and provides a sustainable platform for continued collaboration.

Future Collaboration Value: The COST network established during this project created lasting research partnerships that will continue beyond the funding period. The Deutsche Borse collaboration has produced ongoing data access agreements. The academic partnerships with Humboldt-University Berlin, Bucharest University, and Babes-Bolyai University have resulted in joint PhD supervision arrangements and publication pipelines. These collaborations have positioned the project as a European hub for narrative finance research, with visibility through the COST FinAI Wiki (wiki.fin-ai.eu) and AI-in-Finance.eu platforms."""

print(f"Section 1.4b (COST Action) character count: {len(section_1_4b)} (limit: 4000)")
doc.add_paragraph(section_1_4b)

# Save
output_path = Path(__file__).parent / 'Form_Scientific_Report_IC_ACCURATE.docx'
doc.save(output_path)
print(f"\nReport saved to: {output_path}")

# Summary
print("\n=== CHARACTER COUNT SUMMARY ===")
print(f"Section 1.1: {len(section_1_1):,} / 6,000 chars {'OK' if len(section_1_1) <= 6000 else 'OVER LIMIT'}")
print(f"Section 1.2: {len(section_1_2):,} / 2,000 chars {'OK' if len(section_1_2) <= 2000 else 'OVER LIMIT'}")
print(f"Section 1.3: {len(section_1_3):,} / 2,000 chars {'OK' if len(section_1_3) <= 2000 else 'OVER LIMIT'}")
print(f"Section 1.4a: {len(section_1_4a):,} / 4,000 chars {'OK' if len(section_1_4a) <= 4000 else 'OVER LIMIT'}")
print(f"Section 1.4b: {len(section_1_4b):,} / 4,000 chars {'OK' if len(section_1_4b) <= 4000 else 'OVER LIMIT'}")
