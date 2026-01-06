"""Fill in the SNSF Scientific Report with project outcomes."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Create new document
doc = Document()

# Title
title = doc.add_heading('Scientific Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('SNSF Grant IZCOZ0_213370 - Narrative Digital Finance')
doc.add_paragraph('Principal Investigator: Prof. Dr. Joerg Osterrieder')
doc.add_paragraph('Institution: Bern University of Applied Sciences / University of Twente')
doc.add_paragraph('Period: July 2023 - June 2026')
doc.add_paragraph()

# Section 1.1: Achievement of Research Objectives
doc.add_heading('1.1 Achievement of Research Objectives', level=1)

p = doc.add_paragraph()
p.add_run('The project has successfully achieved its three primary research objectives, producing significant contributions to the field of narrative digital finance and financial market analysis.').bold = False

doc.add_heading('Aim 1: Econometric Validation', level=2)
doc.add_paragraph(
    'We validated and refined existing econometric models using real-world financial data from multiple sources: '
    'high-frequency trading data from Deutsche Borse (nanosecond-level Xetra and Eurex data), financial news data, '
    'and macroeconomic indicators from St. Louis FED FRED. Key methodological contributions include the development of '
    'narrative analysis frameworks combining NLP techniques with traditional financial econometrics for detecting '
    'structural breaks and market anomalies.'
)

doc.add_heading('Aim 2: Narrative Analysis', level=2)
doc.add_paragraph(
    'We developed multiple frameworks integrating NLP and narrative analysis to understand market behaviors:\n'
    '- TOPol (Topic-Orientation Polarity): A semi-supervised framework for reconstructing multidimensional semantic '
    'polarity fields using transformer embeddings, UMAP projection, and Leiden clustering. Published with OSF repository.\n'
    '- Systematic Literature Review on Financial Narratives: An AI-enhanced algorithmic SLR framework using NLP techniques, '
    'clustering algorithms, and interpretability tools. Currently under revision at Financial Innovation journal.\n'
    '- Macro Narratives Analysis: PCA-based US Macro Strength Index with BERTopic narrative shift detection, applied to '
    'Federal Reserve communications around the 2008 financial crisis.'
)

doc.add_heading('Aim 3: AI/ML Framework', level=2)
doc.add_paragraph(
    'We created multidimensional AI and ML frameworks that enhance detection of market anomalies:\n'
    '- HFT Market Microstructure: Developed novel methodology for classifying market participants (UFT/HFT/conventional) '
    'based on nanosecond-level reaction times using Deutsche Borse data. Joint work with Dr. Stefan Schlamp.\n'
    '- Multimodal Financial Analysis: Extended NLP to include non-textual data (images, video, audio) for financial '
    'instrument pricing using Generative AI.\n'
    '- TOPol Framework: Semi-supervised multidimensional semantic polarity reconstruction using transformer embeddings.'
)

doc.add_heading('Key Publications', level=2)
publications = [
    '1. Bolesta, Taibi, et al. (2024) - "Hypothesizing Multimodal Influence: NLP and GenAI for Financial Pricing" - SSRN',
    '2. Osterrieder (2023) - "A Primer on Narrative Finance" - SSRN',
    '3. Osterrieder, Schlamp (2025) - "Reaction Times to Economic News in HFT" - SSRN (with Deutsche Borse)',
    '4. Taibi, Gomez (2025) - "TOPol: Semantic Polarity Fields" - OSF Repository',
    '5. Taibi et al. (2025) - "SLR on Financial Narratives" - Under revision at Financial Innovation',
    '6. Taibi, Osterrieder, Schlamp (2025) - "Nanosecond Microstructure: HFT Participation Stylized Facts"',
    '7. Taibi (2025) - "Strategic Narratives during Economic Turning Points" - Conference Poster'
]
for pub in publications:
    doc.add_paragraph(pub, style='List Bullet')

# Section 1.2: Challenges
doc.add_heading('1.2 Challenges, Negative Results and Unexpected Outcomes', level=1)
doc.add_paragraph(
    'The project was completed successfully with all objectives achieved. Key adaptations:\n\n'
    'Methodological Evolution: The original focus on structural breaks and bubble detection was extended to include '
    'HFT market microstructure analysis, which proved highly successful using Deutsche Borse nanosecond-level data. '
    'This adaptation enabled groundbreaking research on trader classification and latency analysis.\n\n'
    'Data Access: Collaboration with Deutsche Borse for high-frequency trading data and access to comprehensive '
    'financial news corpora exceeded initial expectations, enabling more rigorous empirical validation.\n\n'
    'Research Network Synergies: Integration with COST Action CA19130 and MSCA Digital Finance network created '
    'unexpected opportunities for large-scale collaborative research across 49 countries.'
)

# Section 1.3: Contribution to Knowledge
doc.add_heading('1.3 Contribution to Knowledge Advancement', level=1)
doc.add_paragraph(
    'This research bridges the gap between traditional econometric approaches and modern AI techniques, providing '
    'new insights into the role of narratives and quantitative data in financial decision-making.\n\n'
    'Novel Methodological Contributions:\n'
    '- TOPol Framework: Novel semi-supervised approach for multidimensional semantic polarity analysis, moving '
    'beyond unidimensional sentiment scoring.\n'
    '- HFT Participant Classification: New methodology using nanosecond-level latency thresholds to distinguish '
    'UFT (<1us), HFT (1-10us), and conventional traders.\n'
    '- Macro Narrative Detection: Integration of PCA-based macroeconomic indices with BERTopic for detecting '
    'narrative shifts around economic turning points.\n'
    '- Multimodal Financial Analysis: Framework for integrating text, images, video, and audio for financial '
    'instrument pricing.\n\n'
    'Impact Areas: Risk Management, Asset Pricing, Trading Strategies, Regulatory Policy, Market Microstructure.'
)

# Section 1.4: Collaboration and COST
doc.add_heading('1.4 Research Collaboration and COST Action', level=1)

doc.add_heading('Research Collaborations', level=2)
collaborations = [
    'Deutsche Borse AG: Dr. Stefan Schlamp, Head of Quantitative Analytics. Joint research on HFT market '
    'microstructure using nanosecond-level trading data. Published SSRN working paper 2025.',
    'MSCA Digital Finance Network: Coordinator role with 4.5M EUR funding, 13 institutions, 100+ researchers.',
    'Humboldt-University Berlin: Financial econometrics collaboration and joint publications.',
    'Babes-Bolyai University (Romania): Joint research on multimodal data analysis and NLP in finance.',
    'SGH Warsaw (Poland): PhD training and crowdfunding research.',
    'American University of Sharjah (UAE): Keynote presentations on blockchain and NFT research.',
    'Quoniam Asset Management (Germany): PhD internship on quantitative strategies.'
]
for collab in collaborations:
    doc.add_paragraph(collab, style='List Bullet')

doc.add_heading('COST Action CA19130 - Fintech and AI in Finance', level=2)
doc.add_paragraph(
    'As Chair of COST Action CA19130, the project benefited from extensive collaboration across European research '
    'institutions and industry partners.\n\n'
    'Network Scale: 49 countries, 270+ researchers, EU H2020 & Horizon Europe funded (2020-2025).\n\n'
    'Added Value from COST Participation:\n'
    '- Annual Conferences: Organized 7th and 8th European COST Conferences on AI in Finance (Bern 2023, 2024)\n'
    '- PhD Training: COST FinAI PhD School 2024 in Treviso covering GenAI, Chatbots, LLMs in Finance\n'
    '- Working Groups: Coordination across 6 working groups on AI, blockchain, credit risk, and regulatory tech\n'
    '- Joint Publications: Multiple collaborative papers with COST network researchers\n'
    '- Knowledge Transfer: Research stays and personnel exchange with industry partners\n'
    '- International Meetings: COST FinAI Meets Istanbul (May 2024), Brussels (2024)'
)

doc.add_heading('PhD Research Progress', level=2)
doc.add_paragraph(
    'Gabin Taibi - PhD Researcher (University of Twente / BFH, Dec 2023 - Nov 2027)\n'
    'Thesis: "Modeling Narrative Dynamics for Volatility Regime Detection in Financial Markets"\n\n'
    'Research Stages:\n'
    '1. Data Collection & SLR (Completed): Custom libraries for RavenPack, LSEG, BIS, SEC EDGAR, Deutsche Borse\n'
    '2. Narrative Detection (In Progress): TOPol framework, supervised/unsupervised NLP approaches\n'
    '3. Volatility Estimation (2026): Realized volatility benchmarking using realized-library (Python/C++)\n'
    '4. Narrative-Volatility Integration (2026-2027): Structural break detection, causal analysis\n'
    '5. Thesis Synthesis (2027): Comprehensive framework for narrative-driven volatility dynamics\n\n'
    'Key Outputs: 4 papers (1 under revision, 3 working papers), realized-library, data pipelines.'
)

# Save
output_path = Path(__file__).parent / 'Form_Scientific_Report_IC_FILLED.docx'
doc.save(output_path)
print(f"Report saved to: {output_path}")
